from docx import Document
import sys

path = sys.argv[1]
d = Document(path)

print("===PARAGRAPHS===")
for p in d.paragraphs:
    if p.text.strip():
        style = p.style.name if p.style else ""
        print(f"[{style}] {p.text}")

print("\n===TABLES===")
for i, t in enumerate(d.tables):
    print(f"\n--- Table {i} ---")
    for row in t.rows:
        cells = [c.text.strip() for c in row.cells]
        print(" | ".join(cells))
